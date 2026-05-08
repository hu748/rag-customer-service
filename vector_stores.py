from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
import config_data as config
import os

class VectorStoreService(object):
    def __init__(self, embedding):
        self.embedding = embedding
        self.vector_store = Chroma(
            collection_name=config.collections_name,
            embedding_function=self.embedding,
            persist_directory=config.persist_directory
        )
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.chunk_size,
            chunk_overlap=config.chunk_overlap,
            separators=config.separators,
        )

    def get_retriever(self):
        return self.vector_store.as_retriever(search_kwargs={"k": config.similarity_threshold})

    def add_documents(self, documents):
        if not documents:
            return 0
        
        split_docs = self.text_splitter.split_documents(documents)
        self.vector_store.add_documents(split_docs)
        return len(split_docs)

    def add_text(self, text, metadata=None):
        docs = [Document(page_content=text, metadata=metadata or {})]
        return self.add_documents(docs)

    def _load_file(self, file_path):
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()
        
        if ext == ".pdf":
            return self._load_pdf_langchain(file_path)
        elif ext == ".docx":
            return self._load_docx_langchain(file_path)
        elif ext in [".xlsx", ".xls"]:
            return self._load_excel_langchain(file_path)
        elif ext in [".txt", ".md", ".json"]:
            return self._load_text_file(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

    def _load_pdf_langchain(self, file_path):
        try:
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(file_path)
            docs = loader.load()
            return docs
        except ImportError:
            try:
                from langchain_community.document_loaders import PDFMinerLoader
                loader = PDFMinerLoader(file_path)
                return loader.load()
            except ImportError:
                return self._load_pdf_fallback(file_path)
        except Exception as e:
            return self._load_pdf_fallback(file_path)

    def _load_pdf_fallback(self, file_path):
        text = ""
        try:
            with open(file_path, 'rb') as f:
                import fitz
                doc = fitz.open(file_path)
                for page in doc:
                    text += page.get_text() + "\n"
            return [Document(page_content=text.strip(), metadata={"source": file_path})]
        except ImportError:
            raise ValueError(
                "无法解析PDF文件。请安装以下依赖之一：\n"
                "pip install PyPDF2\n"
                "pip install pdfminer.six\n"
                "pip install pymupdf"
            )
        except Exception as e:
            raise ValueError(f"PDF解析失败: {str(e)}")

    def _load_docx_langchain(self, file_path):
        try:
            from langchain_community.document_loaders import Docx2txtLoader
            loader = Docx2txtLoader(file_path)
            return loader.load()
        except ImportError:
            try:
                from langchain_community.document_loaders import UnstructuredWordDocumentLoader
                loader = UnstructuredWordDocumentLoader(file_path)
                return loader.load()
            except ImportError:
                return self._load_docx_fallback(file_path)

    def _load_docx_fallback(self, file_path):
        try:
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return [Document(page_content=text.strip(), metadata={"source": file_path})]
        except ImportError:
            raise ValueError(
                "无法解析Word文件。请安装以下依赖之一：\n"
                "pip install docx2txt\n"
                "pip install python-docx\n"
                "pip install unstructured"
            )
        except Exception as e:
            raise ValueError(f"Word文档解析失败: {str(e)}")

    def _load_excel_langchain(self, file_path):
        try:
            from langchain_community.document_loaders import UnstructuredExcelLoader
            loader = UnstructuredExcelLoader(file_path)
            return loader.load()
        except ImportError:
            try:
                from langchain_community.document_loaders import ExcelLoader
                loader = ExcelLoader(file_path)
                return loader.load()
            except ImportError:
                return self._load_excel_fallback(file_path)

    def _load_excel_fallback(self, file_path):
        try:
            import pandas as pd
            df = pd.read_excel(file_path)
            text = df.to_string(index=False)
            return [Document(page_content=text.strip(), metadata={"source": file_path})]
        except ImportError:
            raise ValueError(
                "无法解析Excel文件。请安装：\n"
                "pip install pandas openpyxl xlrd"
            )
        except Exception as e:
            raise ValueError(f"Excel文档解析失败: {str(e)}")

    def _load_text_file(self, file_path):
        try:
            from langchain_community.document_loaders import TextLoader
            loader = TextLoader(file_path, encoding='utf-8')
            return loader.load()
        except Exception:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            return [Document(page_content=text, metadata={"source": file_path})]

    def add_file(self, file_path, metadata=None):
        if not os.path.exists(file_path):
            return 0
        
        try:
            docs = self._load_file(file_path)
            
            for doc in docs:
                if metadata:
                    doc.metadata.update(metadata)
                doc.metadata.setdefault('source', file_path)
                doc.metadata.setdefault('file_name', os.path.basename(file_path))
            
            return self.add_documents(docs)
        except Exception as e:
            raise e

    def get_collection_stats(self):
        return self.vector_store.get()

    def clear_all_documents(self):
        self.vector_store.delete_collection()
        self.vector_store = Chroma(
            collection_name=config.collections_name,
            embedding_function=self.embedding,
            persist_directory=config.persist_directory
        )

    @staticmethod
    def get_supported_formats():
        return ["txt", "md", "json", "pdf", "docx", "xlsx", "xls"]

if __name__ == '__main__':
    from langchain_community.embeddings import DashScopeEmbeddings
    retriever = VectorStoreService(DashScopeEmbeddings(model="text-embedding-v4", dashscope_api_key="sk-8f296b06c5cd4d10aa13d8eae463e6a8")).get_retriever()
    
